from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

import fitz
from docx import Document as WordDocument

from app.config import AppConfig
from app.pipeline.identity import (
    DocumentIdentityInput,
    compute_sha256_hex,
)
from app.pipeline.ingestion import (
    DOCXIngestionPayload,
    DOCXLineFragment,
    PDFIngestionPayload,
    PDFLineFragment,
    UnknownDocumentIdentityError,
    build_docx_coordinate_map,
    build_pdf_coordinate_map,
    read_pdf_text_ratios_default,
    route_document,
)


def _fixture_path(file_name: str) -> Path:
    return Path(__file__).parent / "fixtures" / file_name


def _load_fixture(file_name: str) -> dict[str, object]:
    return json.loads(_fixture_path(file_name).read_text(encoding="utf-8"))


def _create_test_pdf(path: Path, lines: list[str]) -> None:
    document = fitz.open()
    page = document.new_page()
    y_offset = 72

    for line in lines:
        page.insert_text((72, y_offset), line)
        y_offset += 24

    document.save(path)
    document.close()


def _create_test_docx(path: Path) -> None:
    document = WordDocument()
    paragraph = document.add_paragraph()
    label = paragraph.add_run("Range:")
    label.bold = True
    paragraph.add_run(" 60 yards")
    document.add_paragraph("Bless")
    document.save(path)


class IngestionPipelineTests(unittest.TestCase):
    def test_default_pdf_ratio_reader_reports_non_zero_text_ratio_for_text_pdf(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = Path(tmp_dir) / "digital-ratio.pdf"
            _create_test_pdf(pdf_path, ["Magic Missile", "Range: 60 yards"])

            text_ratios = read_pdf_text_ratios_default(pdf_path)

            self.assertEqual(len(text_ratios), 1)
            self.assertGreater(text_ratios[0], 0.0)

    def test_default_digital_pdf_backend_extracts_markdown_and_bounding_boxes(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = Path(tmp_dir) / "digital-default.pdf"
            _create_test_pdf(pdf_path, ["Magic Missile", "Range: 60 yards"])
            source_sha = compute_sha256_hex(pdf_path)
            config = AppConfig(document_names_by_sha256={source_sha: "Wizard's Spell Compendium"})

            routed = route_document(
                pdf_path,
                config=config,
                read_pdf_text_ratios=lambda _path: [0.8],
            )

            self.assertEqual(routed.file_type, "pdf")
            self.assertEqual(routed.ingestion_mode, "pdf_digital")
            self.assertIn("Magic Missile", routed.markdown_text)
            self.assertIn("Range: 60 yards", routed.markdown_text)
            self.assertEqual([line for line, _ in routed.coordinate_map.lines], ["Magic Missile", "Range: 60 yards"])
            self.assertEqual(routed.default_source_pages, [1, 1])

            first_region = routed.coordinate_map.get_region(0)
            self.assertIsNotNone(first_region)
            if first_region is None:
                self.fail("Expected PDF region to be present")
            self.assertEqual(first_region.page, 0)
            self.assertIsNotNone(first_region.bbox)

    def test_default_ocr_pdf_backend_uses_tesseract_line_data(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = Path(tmp_dir) / "ocr-default.pdf"
            _create_test_pdf(pdf_path, ["Scanned placeholder"])
            source_sha = compute_sha256_hex(pdf_path)
            config = AppConfig(
                document_names_by_sha256={source_sha: "Scanned Tome"},
                force_ocr_by_sha256={source_sha: True},
            )

            fake_ocr_data = {
                "text": ["Magic", "Missile", "", "Range:", "60", "yards"],
                "left": [10, 60, 0, 10, 70, 95],
                "top": [20, 20, 0, 50, 50, 50],
                "width": [40, 50, 0, 50, 20, 40],
                "height": [10, 10, 0, 10, 10, 10],
                "block_num": [1, 1, 0, 2, 2, 2],
                "par_num": [1, 1, 0, 1, 1, 1],
                "line_num": [1, 1, 0, 1, 1, 1],
            }

            with patch(
                "app.pipeline.ingestion._tesseract_image_to_data",
                return_value=fake_ocr_data,
                create=True,
            ):
                routed = route_document(pdf_path, config=config)

            self.assertEqual(routed.file_type, "pdf")
            self.assertEqual(routed.ingestion_mode, "pdf_ocr")
            self.assertEqual(
                [line for line, _ in routed.coordinate_map.lines],
                ["Magic Missile", "Range: 60 yards"],
            )
            self.assertEqual(routed.default_source_pages, [1, 1])

            first_region = routed.coordinate_map.get_region(0)
            second_region = routed.coordinate_map.get_region(1)
            self.assertIsNotNone(first_region)
            self.assertIsNotNone(second_region)
            if first_region is None or second_region is None:
                self.fail("Expected OCR regions to be present")
            self.assertEqual(first_region.bbox, (10.0, 20.0, 110.0, 30.0))
            self.assertEqual(second_region.bbox, (10.0, 50.0, 135.0, 60.0))

    def test_default_docx_backend_extracts_markdown_and_character_offsets(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "default.docx"
            _create_test_docx(docx_path)
            source_sha = compute_sha256_hex(docx_path)
            config = AppConfig(document_names_by_sha256={source_sha: "Cleric Compendium"})

            routed = route_document(docx_path, config=config)

            first_line = "**Range:** 60 yards"
            second_line = "Bless"

            self.assertEqual(routed.file_type, "docx")
            self.assertEqual(routed.ingestion_mode, "docx")
            self.assertEqual(routed.markdown_text.splitlines(), [first_line, second_line])
            self.assertEqual([line for line, _ in routed.coordinate_map.lines], [first_line, second_line])
            self.assertEqual(routed.default_source_pages, [None, None])

            first_region = routed.coordinate_map.get_region(0)
            second_region = routed.coordinate_map.get_region(1)
            self.assertIsNotNone(first_region)
            self.assertIsNotNone(second_region)
            if first_region is None or second_region is None:
                self.fail("Expected DOCX regions to be present")
            self.assertEqual(first_region.page, -1)
            self.assertEqual(first_region.char_offset, (0, len(first_line)))
            self.assertEqual(
                second_region.char_offset,
                (len(first_line) + 1, len(first_line) + 1 + len(second_line)),
            )

    def test_unsupported_extension_fails_fast_without_side_effects(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            text_path = Path(tmp_dir) / "notes.txt"
            text_path.write_text("plain text", encoding="utf-8")

            config = AppConfig(
                document_names_by_sha256={"a" * 64: "Known Tome"},
                document_offsets={"a" * 64: 4},
                force_ocr_by_sha256={"a" * 64: True},
            )
            expected_names = dict(config.document_names_by_sha256)
            expected_offsets = dict(config.document_offsets)
            expected_force_ocr = dict(config.force_ocr_by_sha256)

            def _resolver(_sha: str) -> DocumentIdentityInput:
                self.fail("Resolver should not be called for unsupported extensions.")

            with patch(
                "app.pipeline.ingestion.compute_sha256_hex",
                side_effect=AssertionError(
                    "SHA should not be computed for unsupported extensions."
                ),
            ):
                with self.assertRaises(ValueError) as exc_info:
                    route_document(
                        text_path,
                        config=config,
                        resolve_unknown_identity=_resolver,
                    )

            self.assertEqual(
                str(exc_info.exception),
                "Unsupported file extension '.txt'. Expected .pdf or .docx.",
            )
            self.assertEqual(config.document_names_by_sha256, expected_names)
            self.assertEqual(config.document_offsets, expected_offsets)
            self.assertEqual(config.force_ocr_by_sha256, expected_force_ocr)

    def test_unknown_document_hash_requires_identity_metadata_before_routing(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = Path(tmp_dir) / "sample.pdf"
            pdf_path.write_bytes(b"pdf bytes")

            with self.assertRaises(UnknownDocumentIdentityError):
                route_document(
                    pdf_path,
                    config=AppConfig(),
                    read_pdf_text_ratios=lambda _path: [0.4],
                    ingest_pdf_digital=lambda _path: PDFIngestionPayload(
                        markdown_text="ignored",
                        lines=[],
                    ),
                )

    def test_unknown_sha_resolver_with_none_display_name_raises_domain_error(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = Path(tmp_dir) / "bad-identity.pdf"
            pdf_path.write_bytes(b"pdf bytes")

            with self.assertRaises(UnknownDocumentIdentityError):
                route_document(
                    pdf_path,
                    config=AppConfig(),
                    resolve_unknown_identity=lambda _sha: DocumentIdentityInput(
                        source_display_name=None,
                        page_offset=0,
                        force_ocr=False,
                    ),
                )

    def test_unknown_sha_resolver_with_invalid_page_offset_raises_domain_error(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = Path(tmp_dir) / "bad-offset.pdf"
            pdf_path.write_bytes(b"pdf bytes")

            with self.assertRaises(UnknownDocumentIdentityError):
                route_document(
                    pdf_path,
                    config=AppConfig(),
                    resolve_unknown_identity=lambda _sha: DocumentIdentityInput(
                        source_display_name="Arcane Reference",
                        page_offset="not-an-int",
                        force_ocr=False,
                    ),
                )

    def test_unknown_sha_resolver_with_bool_page_offset_raises_domain_error(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = Path(tmp_dir) / "bool-offset.pdf"
            pdf_path.write_bytes(b"pdf bytes")

            with self.assertRaises(UnknownDocumentIdentityError):
                route_document(
                    pdf_path,
                    config=AppConfig(),
                    resolve_unknown_identity=lambda _sha: DocumentIdentityInput(
                        source_display_name="Arcane Reference",
                        page_offset=True,
                        force_ocr=False,
                    ),
                    read_pdf_text_ratios=lambda _path: [0.8],
                    ingest_pdf_digital=lambda _path: PDFIngestionPayload(
                        markdown_text="digital route",
                        lines=[],
                    ),
                )

    def test_unknown_sha_resolver_with_fractional_numeric_page_offset_raises_domain_error(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = Path(tmp_dir) / "fractional-offset.pdf"
            pdf_path.write_bytes(b"pdf bytes")

            with self.assertRaises(UnknownDocumentIdentityError):
                route_document(
                    pdf_path,
                    config=AppConfig(),
                    resolve_unknown_identity=lambda _sha: DocumentIdentityInput(
                        source_display_name="Arcane Reference",
                        page_offset=1.9,
                        force_ocr=False,
                    ),
                    read_pdf_text_ratios=lambda _path: [0.8],
                    ingest_pdf_digital=lambda _path: PDFIngestionPayload(
                        markdown_text="digital route",
                        lines=[],
                    ),
                )

    def test_unknown_sha_resolver_with_string_true_force_ocr_routes_to_ocr(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = Path(tmp_dir) / "string-true-force-ocr.pdf"
            pdf_path.write_bytes(b"pdf bytes")
            source_sha = compute_sha256_hex(pdf_path)
            config = AppConfig()

            routed = route_document(
                pdf_path,
                config=config,
                resolve_unknown_identity=lambda _sha: {
                    "source_display_name": "Arcane Reference",
                    "page_offset": 0,
                    "force_ocr": "true",
                },
                ingest_pdf_digital=lambda _path: self.fail(
                    "Digital ingestor should not be called when force_ocr is true."
                ),
                ingest_pdf_ocr=lambda _path: PDFIngestionPayload(
                    markdown_text="ocr route",
                    lines=[
                        PDFLineFragment(
                            text="ocr line",
                            page=0,
                            bbox=(0.0, 0.0, 10.0, 10.0),
                        )
                    ],
                ),
            )

            self.assertTrue(routed.identity.force_ocr)
            self.assertEqual(routed.ingestion_mode, "pdf_ocr")
            self.assertTrue(config.force_ocr_by_sha256[source_sha])

    def test_unknown_sha_resolver_with_string_false_force_ocr_stays_digital(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = Path(tmp_dir) / "string-false-force-ocr.pdf"
            pdf_path.write_bytes(b"pdf bytes")
            source_sha = compute_sha256_hex(pdf_path)
            config = AppConfig()

            routed = route_document(
                pdf_path,
                config=config,
                resolve_unknown_identity=lambda _sha: {
                    "source_display_name": "Arcane Reference",
                    "page_offset": 0,
                    "force_ocr": "false",
                },
                read_pdf_text_ratios=lambda _path: [0.8],
                ingest_pdf_digital=lambda _path: PDFIngestionPayload(
                    markdown_text="digital route",
                    lines=[
                        PDFLineFragment(
                            text="digital line",
                            page=0,
                            bbox=(0.0, 0.0, 10.0, 10.0),
                        )
                    ],
                ),
                ingest_pdf_ocr=lambda _path: self.fail(
                    "OCR ingestor should not be called when force_ocr is false."
                ),
            )

            self.assertFalse(routed.identity.force_ocr)
            self.assertEqual(routed.ingestion_mode, "pdf_digital")
            self.assertFalse(config.force_ocr_by_sha256[source_sha])

    def test_unknown_sha_resolver_with_invalid_force_ocr_raises_domain_error(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = Path(tmp_dir) / "invalid-force-ocr.pdf"
            pdf_path.write_bytes(b"pdf bytes")
            config = AppConfig()

            with self.assertRaises(UnknownDocumentIdentityError):
                route_document(
                    pdf_path,
                    config=config,
                    resolve_unknown_identity=lambda _sha: {
                        "source_display_name": "Arcane Reference",
                        "page_offset": 0,
                        "force_ocr": "not-a-bool",
                    },
                )

            self.assertEqual(config.force_ocr_by_sha256, {})

    def test_known_sha_reuses_stored_identity_and_force_ocr_override(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = Path(tmp_dir) / "wizard.pdf"
            pdf_path.write_bytes(b"known pdf")
            source_sha = compute_sha256_hex(pdf_path)
            config = AppConfig(
                document_names_by_sha256={source_sha: "Player's Handbook"},
                document_offsets={source_sha: 3},
                force_ocr_by_sha256={source_sha: True},
            )

            fixture = _load_fixture("pdf_coordinate_fixture.json")
            pdf_lines = [
                PDFLineFragment(
                    text=item["text"],
                    page=item["page"],
                    bbox=tuple(item["bbox"]),
                )
                for item in fixture["lines"]
            ]

            routed = route_document(
                pdf_path,
                config=config,
                read_pdf_text_ratios=lambda _path: [0.9],
                ingest_pdf_digital=lambda _path: PDFIngestionPayload(
                    markdown_text="digital-branch",
                    lines=[],
                ),
                ingest_pdf_ocr=lambda _path: PDFIngestionPayload(
                    markdown_text=fixture["markdown_text"],
                    lines=pdf_lines,
                ),
            )

            self.assertEqual(routed.identity.source_display_name, "Player's Handbook")
            self.assertEqual(routed.identity.page_offset, 3)
            self.assertTrue(routed.identity.force_ocr)
            self.assertEqual(routed.ingestion_mode, "pdf_ocr")
            self.assertEqual(routed.default_source_pages, [4, 4])

    def test_force_ocr_routes_to_ocr_without_ratio_reader_adapter(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = Path(tmp_dir) / "forced-ocr.pdf"
            pdf_path.write_bytes(b"forced ocr")
            source_sha = compute_sha256_hex(pdf_path)
            config = AppConfig(
                document_names_by_sha256={source_sha: "Forced OCR Tome"},
                force_ocr_by_sha256={source_sha: True},
            )

            routed = route_document(
                pdf_path,
                config=config,
                ingest_pdf_digital=lambda _path: self.fail("Digital ingestor should not be called"),
                ingest_pdf_ocr=lambda _path: PDFIngestionPayload(
                    markdown_text="ocr route",
                    lines=[
                        PDFLineFragment(
                            text="ocr line",
                            page=0,
                            bbox=(0.0, 0.0, 10.0, 10.0),
                        )
                    ],
                ),
            )

            self.assertEqual(routed.ingestion_mode, "pdf_ocr")
            self.assertEqual(routed.markdown_text, "ocr route")

    def test_non_forced_sparse_pdf_uses_default_ratio_reader_to_route_to_ocr(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = Path(tmp_dir) / "regular.pdf"
            _create_test_pdf(pdf_path, ["regular pdf"])
            source_sha = compute_sha256_hex(pdf_path)
            config = AppConfig(document_names_by_sha256={source_sha: "Regular Tome"})

            routed = route_document(
                pdf_path,
                config=config,
                ingest_pdf_digital=lambda _path: self.fail(
                    "Sparse page should route to OCR when the default ratio reader is used."
                ),
                ingest_pdf_ocr=lambda _path: PDFIngestionPayload(
                    markdown_text="ocr route",
                    lines=[
                        PDFLineFragment(
                            text="ocr route",
                            page=0,
                            bbox=(0.0, 0.0, 10.0, 10.0),
                        )
                    ],
                )
            )

            self.assertEqual(routed.ingestion_mode, "pdf_ocr")
            self.assertEqual(routed.markdown_text, "ocr route")
    def test_pdf_source_page_defaults_use_document_offset(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = Path(tmp_dir) / "offsets.pdf"
            pdf_path.write_bytes(b"offset pdf")
            source_sha = compute_sha256_hex(pdf_path)
            config = AppConfig(
                document_names_by_sha256={source_sha: "Offset Tome"},
                document_offsets={source_sha: 5},
            )

            routed = route_document(
                pdf_path,
                config=config,
                read_pdf_text_ratios=lambda _path: [0.8],
                ingest_pdf_digital=lambda _path: PDFIngestionPayload(
                    markdown_text="line a\nline b",
                    lines=[
                        PDFLineFragment(
                            text="line a",
                            page=0,
                            bbox=(10.0, 10.0, 50.0, 20.0),
                        ),
                        PDFLineFragment(
                            text="line b",
                            page=2,
                            bbox=(10.0, 30.0, 50.0, 40.0),
                        ),
                    ],
                ),
            )

            self.assertEqual(routed.default_source_pages, [6, 8])

    def test_docx_with_inconsistent_page_sequence_leaves_source_pages_unset(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "spells.docx"
            docx_path.write_bytes(b"docx bytes")
            source_sha = compute_sha256_hex(docx_path)
            config = AppConfig(document_names_by_sha256={source_sha: "Cleric Compendium"})

            fixture = _load_fixture("docx_coordinate_fixture.json")
            docx_lines = [
                DOCXLineFragment(
                    text=item["text"],
                    char_offset=tuple(item["char_offset"]),
                )
                for item in fixture["lines"]
            ]

            routed = route_document(
                docx_path,
                config=config,
                ingest_docx=lambda _path: DOCXIngestionPayload(
                    markdown_text=fixture["markdown_text"],
                    lines=docx_lines,
                    page_sequence=fixture["page_sequence"],
                ),
            )

            self.assertEqual(routed.file_type, "docx")
            self.assertEqual(routed.default_source_pages, [None, None])

    def test_docx_with_non_monotonic_page_sequence_and_matching_line_count_leaves_source_pages_unset(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "spells.docx"
            docx_path.write_bytes(b"docx bytes")
            source_sha = compute_sha256_hex(docx_path)
            config = AppConfig(document_names_by_sha256={source_sha: "Cleric Compendium"})

            docx_lines = [
                DOCXLineFragment(text="Line 1", char_offset=(0, 6)),
                DOCXLineFragment(text="Line 2", char_offset=(7, 13)),
                DOCXLineFragment(text="Line 3", char_offset=(14, 20)),
            ]

            routed = route_document(
                docx_path,
                config=config,
                ingest_docx=lambda _path: DOCXIngestionPayload(
                    markdown_text="Line 1\nLine 2\nLine 3\n",
                    lines=docx_lines,
                    page_sequence=[1, 3, 2],
                ),
            )

            self.assertEqual(routed.file_type, "docx")
            self.assertEqual(routed.default_source_pages, [None, None, None])

    def test_docx_with_consistent_page_sequence_applies_offset_to_source_pages(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "spells.docx"
            docx_path.write_bytes(b"docx bytes")
            source_sha = compute_sha256_hex(docx_path)
            config = AppConfig(
                document_names_by_sha256={source_sha: "Cleric Compendium"},
                document_offsets={source_sha: 4},
            )

            fixture = _load_fixture("docx_coordinate_fixture.json")
            docx_lines = [
                DOCXLineFragment(
                    text=item["text"],
                    char_offset=tuple(item["char_offset"]),
                )
                for item in fixture["lines"]
            ]

            routed = route_document(
                docx_path,
                config=config,
                ingest_docx=lambda _path: DOCXIngestionPayload(
                    markdown_text=fixture["markdown_text"],
                    lines=docx_lines,
                    page_sequence=[1, 3],
                ),
            )

            self.assertEqual(routed.file_type, "docx")
            self.assertEqual(routed.default_source_pages, [5, 7])

    def test_docx_with_fractional_float_page_sequence_leaves_source_pages_unset(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "spells.docx"
            docx_path.write_bytes(b"docx bytes")
            source_sha = compute_sha256_hex(docx_path)
            config = AppConfig(document_names_by_sha256={source_sha: "Cleric Compendium"})

            fixture = _load_fixture("docx_coordinate_fixture.json")
            docx_lines = [
                DOCXLineFragment(
                    text=item["text"],
                    char_offset=tuple(item["char_offset"]),
                )
                for item in fixture["lines"]
            ]

            routed = route_document(
                docx_path,
                config=config,
                ingest_docx=lambda _path: DOCXIngestionPayload(
                    markdown_text=fixture["markdown_text"],
                    lines=docx_lines,
                    page_sequence=[1.9, 2],
                ),
            )

            self.assertEqual(routed.file_type, "docx")
            self.assertEqual(routed.default_source_pages, [None, None])

    def test_docx_with_fractional_string_page_sequence_leaves_source_pages_unset(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "spells.docx"
            docx_path.write_bytes(b"docx bytes")
            source_sha = compute_sha256_hex(docx_path)
            config = AppConfig(document_names_by_sha256={source_sha: "Cleric Compendium"})

            fixture = _load_fixture("docx_coordinate_fixture.json")
            docx_lines = [
                DOCXLineFragment(
                    text=item["text"],
                    char_offset=tuple(item["char_offset"]),
                )
                for item in fixture["lines"]
            ]

            routed = route_document(
                docx_path,
                config=config,
                ingest_docx=lambda _path: DOCXIngestionPayload(
                    markdown_text=fixture["markdown_text"],
                    lines=docx_lines,
                    page_sequence=["1.9", "2"],
                ),
            )

            self.assertEqual(routed.file_type, "docx")
            self.assertEqual(routed.default_source_pages, [None, None])

    def test_unknown_sha_resolver_persists_identity_metadata(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = Path(tmp_dir) / "new-book.pdf"
            pdf_path.write_bytes(b"new book")
            source_sha = compute_sha256_hex(pdf_path)
            config = AppConfig()

            routed = route_document(
                pdf_path,
                config=config,
                resolve_unknown_identity=lambda _sha: DocumentIdentityInput(
                    source_display_name="Wizard's Spell Compendium",
                    page_offset=2,
                    force_ocr=False,
                ),
                read_pdf_text_ratios=lambda _path: [0.6],
                ingest_pdf_digital=lambda _path: PDFIngestionPayload(
                    markdown_text="content",
                    lines=[
                        PDFLineFragment(
                            text="content",
                            page=0,
                            bbox=(0.0, 0.0, 20.0, 10.0),
                        )
                    ],
                ),
            )

            self.assertEqual(routed.identity.source_display_name, "Wizard's Spell Compendium")
            self.assertEqual(config.document_names_by_sha256[source_sha], "Wizard's Spell Compendium")
            self.assertEqual(config.document_offsets[source_sha], 2)
            self.assertFalse(config.force_ocr_by_sha256[source_sha])


class CoordinateMapFixtureTests(unittest.TestCase):
    def test_pdf_coordinate_map_generation_from_fixture(self) -> None:
        fixture = _load_fixture("pdf_coordinate_fixture.json")
        line_fragments = [
            PDFLineFragment(
                text=item["text"],
                page=item["page"],
                bbox=tuple(item["bbox"]),
            )
            for item in fixture["lines"]
        ]

        coordinate_map = build_pdf_coordinate_map(line_fragments)

        self.assertEqual(len(coordinate_map.lines), 2)
        self.assertEqual(coordinate_map.get_line(0), "Magic Missile")
        first_region = coordinate_map.get_region(0)
        self.assertIsNotNone(first_region)
        if first_region is None:
            self.fail("Expected PDF region to be present")
        self.assertEqual(first_region.page, 0)
        self.assertEqual(first_region.bbox, (10.0, 20.0, 120.0, 32.0))

    def test_pdf_coordinate_map_rejects_non_integral_page_values(self) -> None:
        for page_value in (1.9, -0.2):
            with self.subTest(page_value=page_value):
                with self.assertRaisesRegex(ValueError, "integral"):
                    build_pdf_coordinate_map(
                        [
                            PDFLineFragment(
                                text="Magic Missile",
                                page=page_value,
                                bbox=(10.0, 20.0, 120.0, 32.0),
                            )
                        ]
                    )

    def test_docx_coordinate_map_generation_from_fixture(self) -> None:
        fixture = _load_fixture("docx_coordinate_fixture.json")
        line_fragments = [
            DOCXLineFragment(
                text=item["text"],
                char_offset=tuple(item["char_offset"]),
            )
            for item in fixture["lines"]
        ]

        coordinate_map = build_docx_coordinate_map(line_fragments)

        self.assertEqual(len(coordinate_map.lines), 2)
        self.assertEqual(coordinate_map.get_line(1), "Duration: 6 rounds")
        second_region = coordinate_map.get_region(1)
        self.assertIsNotNone(second_region)
        if second_region is None:
            self.fail("Expected DOCX region to be present")
        self.assertEqual(second_region.page, -1)
        self.assertEqual(second_region.char_offset, (6, 24))


if __name__ == "__main__":
    unittest.main()
